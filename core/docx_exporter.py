
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import os

def convert_markdown_to_docx(markdown_text, output_path):
    """
    将Markdown文本转换为Word文档（学术论文标准排版）
    
    排版标准：
    - 一级标题：三号（16pt）加粗，居中
    - 二级标题：小三号（15pt）加粗，左对齐
    - 三级标题：四号（14pt）加粗，左对齐
    - 正文：小四号（12pt），中文宋体，英文Times New Roman
    - 行距：1.5倍
    """
    doc = Document()
    
    # --- 基础样式设置 ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置行距 1.5倍
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # --- 标题样式定义 ---
    # Python-docx的默认Heading样式需要手动调整
    
    def set_heading_style(level, font_size_pt, is_bold=True, alignment=None):
        style_name = f'Heading {level}'
        if style_name not in doc.styles:
            return
            
        style = doc.styles[style_name]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(font_size_pt)
        font.bold = is_bold
        font.color.rgb = RGBColor(0, 0, 0)  # 纯黑
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        if alignment is not None:
            style.paragraph_format.alignment = alignment
            
        # 段前段后间距
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 应用样式配置
    set_heading_style(1, 16, True, WD_ALIGN_PARAGRAPH.CENTER) # 一级：三号(16pt)
    set_heading_style(2, 15, True, WD_ALIGN_PARAGRAPH.LEFT)   # 二级：小三(15pt)
    set_heading_style(3, 14, True, WD_ALIGN_PARAGRAPH.LEFT)   # 三级：四号(14pt)

    # --- 内容解析与写入 ---
    # --- 内容解析与写入 (智能段落合并) ---
    lines = markdown_text.split('\n')
    
    current_paragraph_lines = []
    
    def flush_paragraph():
        if current_paragraph_lines:
            content = ''.join(current_paragraph_lines).strip()
            if content:
                # 清理Markdown加粗符号（如果需要）
                clean_content = content.replace('**', '')
                p = doc.add_paragraph(clean_content)
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            current_paragraph_lines.clear()

    for line in lines:
        stripped_line = line.strip()
        
        # 1. 空行 -> 结束当前段落
        if not stripped_line:
            flush_paragraph()
            continue
            
        # 2. 标题 -> 结束当前段落并写入标题
        if stripped_line.startswith('#'):
            flush_paragraph()
            
            if stripped_line.startswith('### '):
                doc.add_heading(stripped_line.replace('###', '').strip(), level=3)
            elif stripped_line.startswith('## '):
                doc.add_heading(stripped_line.replace('##', '').strip(), level=2)
            elif stripped_line.startswith('# '):
                p = doc.add_heading(stripped_line.replace('#', '').strip(), level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # [*] 2.5 检测中文编号标题（可能没有###前缀）
        # 例如：（一）概念界定、（二）理论框架
        chinese_heading_match = re.match(r'^[（(]([一二三四五六七八九十]+)[）)](.+)$', stripped_line)
        if chinese_heading_match:
            flush_paragraph()
            heading_text = f"（{chinese_heading_match.group(1)}）{chinese_heading_match.group(2).strip()}"
            doc.add_heading(heading_text, level=3)
            continue
            
        # 3. 列表 -> 结束当前段落并写入列表项
        if stripped_line.startswith('- ') or stripped_line.startswith('* '):
            flush_paragraph()
            doc.add_paragraph(stripped_line[2:], style='List Bullet')
            continue
            
        # 4. 普通文本 -> 添加到当前段落缓冲（如果是中文，直接连接；英文可能需要空格，简单起见直接连接）
        # 考虑到“句子1：”这种可能导致了换行，但我们希望连成段
        # 如果前一行以句号/问号/感叹号结尾，可能不需要空格（中文）
        # 这里简单直接拼接，不加换行符，从而把它们合并成一段
        current_paragraph_lines.append(stripped_line)

    # 处理最后一个可能未写入的段落
    flush_paragraph()

    doc.save(output_path)
    return output_path
